from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("reports/qwen3vl_fundus_finetune_plan_report.docx")
STATS = Path("data/fundus_validated/validated.stats.json")
CLEAN_STATS = Path("data/fundus_validated/validated_clean.stats.json")

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
PALE_GREEN = "E2F0D9"
PALE_GRAY = "F2F2F2"
PALE_YELLOW = "FFF2CC"


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, size=8.5, bold=False, fill=None, color=None, align="center"):
    cell.text = ""
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None, font_size=8.5, long=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, size=font_size, bold=True, fill=BLUE, color=(255, 255, 255))
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            fill = PALE_BLUE if r % 2 == 0 else None
            cell_text(cells[i], value, size=font_size, fill=fill, align="left" if long else "center")
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return t


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=15 if level == 1 else 12, bold=True, color=(31, 78, 121))
    return p


def para(doc, text="", size=10, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=size, bold=True)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def note_box(doc, title, body, fill=PALE_YELLOW):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=(31, 78, 121))
    p.add_run("\n")
    r = p.add_run(body)
    set_font(r, size=9.5)
    doc.add_paragraph()


def code_block(doc, text, fill="F7F7F7"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(7.4)
    doc.add_paragraph()


def cot_card(doc, title, meta, system, user, assistant, json_line):
    heading(doc, title, 3)
    para(doc, meta, size=8.8, bold_prefix="真实来源：")
    obj = {
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
            {"role": "assistant", "content": assistant + "\n\n【JSON】\n" + json_line},
        ],
        "images": [meta.split("；")[1] if "；" in meta else ""],
    }
    code_block(doc, json.dumps(obj, ensure_ascii=False, indent=2), fill=PALE_BLUE)


def setup_doc(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.35)
    sec.bottom_margin = Cm(1.25)
    sec.left_margin = Cm(1.35)
    sec.right_margin = Cm(1.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)


def build_doc():
    stats = load_json(STATS)
    clean = load_json(CLEAN_STATS)
    datasets = stats["datasets"]
    usable = clean["usable_totals"]
    present = clean["present_totals"]

    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Qwen3-VL-8B 眼底图像微调方案汇报")
    set_font(run, size=20, bold=True, color=(31, 78, 121))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("基于真实数据统计、RetSAM 量化伪标签、强标注 mask 与分层 CoT/SFT 的训练设计（2026-04-29 更新）")
    set_font(run, size=10)

    heading(doc, "一、方案概览", 1)
    para(
        doc,
        "本方案面向 Qwen3-VL-8B 的眼底图像微调，目标不是直接让模型记忆 DR 等级，而是先建立解剖结构感知和病灶证据识别，再学习基于证据的 DR 分级。"
        "当前 pipeline 已将 APTOS、DDR Grading、IDRiD、FGADR Seg-set 和 DDR lesion_segmentation 汇总为统一事实层，并在清洗后用于生成 L2-L4 分层 CoT/SFT 数据。",
    )
    note_box(
        doc,
        "当前关键数据量",
        f"统一索引图像 {stats['n_records']} 张；RetSAM 有效输出 {usable['L2']} 张；清洗后 L3 可用 {usable['L3']} 张，其中强标注 L3 {usable['strong_L3']} 张、RetSAM L3 {usable['retsam_L3']} 张；L4 可用 {usable['L4']} 张。"
        f"清洗后阳性病灶统计：HE {present['HE']}，EX {present['EX']}，SE {present['SE']}，MA {present['MA']}，IRMA {present['IRMA']}，NV {present['NV']}。",
    )

    heading(doc, "二、RetSAM 与强标注覆盖情况", 1)
    para(doc, "下表保留原 MD 的核心数据结构，用于向导师展示：每个数据集上有哪些标签、哪些来自强标注、哪些来自 RetSAM 伪标签。")
    table(
        doc,
        ["数据集", "分级标注", "病灶标注（Mask/量化）", "eye_side", "CDR", "A/V ratio", "CRAE/CRVE", "tortuosity", "fractal_dimension", "vessel_qc_flag", "OD/黄斑坐标", "备注"],
        [
            ["APTOS", "强标注", "伪标签（RetSAM HE/EX/SE 量化）", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM 仅对 Grade 1-4 计算，Grade 0 使用分级模板。"],
            ["DDR Grading", "强标注", "伪标签（RetSAM HE/EX/SE 量化）", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM；已知 DDR macula 坐标尺度异常", "RetSAM 覆盖 Grade 1-4，Grade 5 已丢弃。"],
            ["IDRiD（train+test）", "强标注", "强标注像素级 mask + RetSAM 量化补充", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "官方定位 + RetSAM", "强 mask 质量高；RetSAM 当前覆盖 train 413 张。"],
            ["FGADR Seg-set", "强标注（官方 CSV）", "强标注像素级 mask + RetSAM 量化补充", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "强 mask 质量高；RetSAM 已覆盖 1842 张。"],
            ["DDR lesion_segmentation", "无", "强标注像素级 mask + RetSAM 量化补充", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "RetSAM", "分割集不带 grade；RetSAM 已覆盖 train/valid/test 共 757 张。"],
        ],
        widths=[2.4, 2.2, 4.1, 1.5, 1.2, 1.5, 1.7, 1.7, 2.1, 1.9, 2.9, 4.0],
        font_size=6.7,
        long=True,
    )
    table(
        doc,
        ["数据集", "索引图像数", "RetSAM 有效数", "强标注结构化数", "备注"],
        [
            ["APTOS", datasets["aptos::all"]["n"], datasets["aptos::all"].get("retsam_present", 0), 0, "RetSAM 覆盖 Grade 1-4"],
            ["DDR Grading", datasets["ddr_grading::all"]["n"], datasets["ddr_grading::all"].get("retsam_present", 0), 0, "Grade 5 已丢弃"],
            ["IDRiD train", datasets["idrid::train"]["n"], datasets["idrid::train"].get("retsam_present", 0), datasets["idrid::train"].get("strong_annotation_present", 0), "强 mask + RetSAM"],
            ["IDRiD test", datasets["idrid::test"]["n"], datasets["idrid::test"].get("retsam_present", 0), datasets["idrid::test"].get("strong_annotation_present", 0), "保留为评估，不进入训练"],
            ["FGADR Seg-set", datasets["fgadr_seg::all"]["n"], datasets["fgadr_seg::all"].get("retsam_present", 0), datasets["fgadr_seg::all"].get("strong_annotation_present", 0), "强 mask + RetSAM"],
            ["DDR lesion train", datasets["ddr_seg::train"]["n"], datasets["ddr_seg::train"].get("retsam_present", 0), 0, "L3 可用"],
            ["DDR lesion valid", datasets["ddr_seg::valid"]["n"], datasets["ddr_seg::valid"].get("retsam_present", 0), 0, "已补齐 RetSAM"],
            ["DDR lesion test", datasets["ddr_seg::test"]["n"], datasets["ddr_seg::test"].get("retsam_present", 0), 0, "已补齐 RetSAM"],
        ],
        widths=[4.0, 2.2, 2.2, 2.4, 7.0],
        font_size=8.2,
    )

    heading(doc, "三、清洗后的可信事实层", 1)
    para(
        doc,
        "CoT 生成只读取 validated_clean.jsonl，不直接读取原始 RetSAM JSON。清洗逻辑遵循“强标注优先、RetSAM 次之、规则模板兜底”的顺序；低置信或低 QC 字段不强行生成结论，而是作为 unknown/拒答样本训练模型。",
    )
    table(
        doc,
        ["字段", "当前处理", "训练含义"],
        [
            ["MA", "RetSAM 不提供 MA；阳性来自 IDRiD/FGADR 强 mask，Grade 1 可生成 template_only", "不能写成 RetSAM 检出 MA"],
            ["HE/EX", "RetSAM 阳性经过面积、数量和置信度过滤；强 mask 优先", "用于显式病灶描述和 L4 证据绑定"],
            ["SE", "清洗规则最保守；低置信或小灶降级为 false/unknown", "重点训练不编造软性渗出"],
            ["eye_side/CDR", "OD 模块 QC 通过时保留；坐标异常不自动否定 eye_side", "用于 L2 左右眼和 CDR 学习"],
            ["血管指标", "vessel QC 失败时 A/V、tortuosity、CRAE/CRVE 置 unknown", "训练模型在不可靠时拒答"],
        ],
        widths=[3.0, 9.0, 7.0],
        font_size=8.4,
        long=True,
    )

    heading(doc, "四、CoT 问题分类设计", 1)
    para(
        doc,
        "问题构建采用 FunBench 的分层思想，但结合本项目真实字段重新落地。每条样本只训练一个清晰能力点：L2 关注解剖结构，L3 关注病灶属性和负担，L4 关注证据绑定诊断。"
        "system 负责规则边界，user 只给观察任务，assistant 才给显式视觉描述、结构化证据和结论。",
    )
    table(
        doc,
        ["层级", "子任务", "生成条件", "用户问题方向", "答案学习点"],
        [
            ["L2-1", "laterality", "eye_side.valid=true", "判断左眼/右眼", "描述视盘与黄斑相对关系，再输出 eye_side"],
            ["L2-2", "CDR", "cdr.valid=true", "观察视盘和视杯，估计杯盘比", "说明视盘、视杯和垂直径比例，输出 cdr_bucket"],
            ["L2-3", "vessel/QC", "血管指标有效或 QC 失败", "判断 A/V 与迂曲度是否可靠", "QC 失败时输出 unknown，不硬猜"],
            ["L3-1/2/3", "MA/HE/EX 显式识别", "强 mask 或清洗后阳性", "观察红点、暗红斑块、亮黄色沉积", "颜色、形态、边界、数量到病灶名的映射"],
            ["L3-4", "SE 识别/拒答", "SE 阳性或被清洗降级", "观察灰白棉絮样病灶", "低置信 SE 不作为阳性证据"],
            ["L3-5/6", "lesion-only 与负担", "病灶 presence/count/area 可用", "选择可见病灶并描述负担", "输出 lesions、count_bucket、area_bucket"],
            ["L4-1/2", "Grade 0/1", "grade=0 或 grade=1", "核查是否有可靠 DR 病灶证据", "Grade 1 区分 template_only MA 与强标注 MA"],
            ["L4-3/4", "证据绑定分级/冲突核查", "grade>=2 且 L3 证据可用，或存在冲突", "先核查病灶证据，再判断分级是否有依据", "DR grade 必须引用 L3 证据，冲突时提示复核"],
        ],
        widths=[2.0, 3.3, 4.2, 5.6, 7.4],
        font_size=8.0,
        long=True,
    )

    heading(doc, "五、真实 CoT 样例", 1)
    cot_card(
        doc,
        "样例 1：L2 laterality 与 CDR",
        "aptos::all::000c1434d8d7；data/cropped/aptos/grade1_4/000c1434d8d7.png；eye_side=right，CDR=0.3495，均来自 validated_retsam。",
        "你是眼底图像分析助手。只回答指定 L2 任务；先说明视觉依据，再给结构化结论。",
        "<image> 请判断这张眼底图来自左眼还是右眼，并估计杯盘比是否偏高。",
        "【观察】laterality 依据视盘与黄斑/中央凹的相对关系：视盘位于鼻侧，黄斑位于颞侧。本图事实层给出 eye_side=right。视盘为较亮橙黄色圆/椭圆结构，视杯位于视盘中央且颜色更浅，CDR 约为 0.35，未见明显杯盘比增大。\n【结论】右眼；CDR 属正常或轻度范围。",
        "{\"task\":\"L2_anatomy\",\"eye_side\":\"right\",\"cdr\":0.3495,\"cdr_bucket\":\"normal_or_mild\",\"source\":\"validated_retsam\"}",
    )
    cot_card(
        doc,
        "样例 2：L3 强标注显式病灶",
        "idrid::train::IDRiD_001；data/idrid/images/train/IDRiD_001.jpg；MA count=18，HE count=15，EX area=12027.0，source=strong_mask_stage1_easy。",
        "你是眼底病灶识别助手。只做病灶分析，不输出 DR grade；先描述颜色、形态、边界和数量，再映射到病灶名。",
        "<image> 请观察是否存在微小红色圆点、暗红不规则斑块、亮黄色边界清楚沉积或灰白棉絮样病灶。",
        "【观察】可见多处微小红色规则圆点，另有多处暗红不规则斑块状病灶；同时可见较大面积、亮黄色、边界较清楚的多灶散在沉积。未见明确灰白棉絮样病灶或新生血管。\n【证据】MA present=true, count=18；HE present=true, count=15；EX present=true, area=12027.0；SE/NV present=false。\n【结论】可见 MA、HE、EX；未见可靠 SE 或 NV。",
        "{\"task\":\"L3_explicit_lesion_detection\",\"lesions\":[\"MA\",\"HE\",\"EX\"],\"absent\":[\"SE\",\"NV\"],\"source\":\"strong_mask_stage1_easy\"}",
    )
    cot_card(
        doc,
        "样例 3：L4 Grade 1 模板，不伪造 MA",
        "aptos::all::0024cdab0c1e；data/cropped/aptos/grade1_4/0024cdab0c1e.png；grade=1，MA=template_only，HE/EX/SE=false。",
        "你是眼底分级助手。RetSAM 不提供 MA；template_only 不能写成图像直接检出，也不能写成 RetSAM 检出。",
        "<image> 该图标注为轻度 DR，应如何给出证据解释？",
        "【观察】清洗后未保留可靠 HE、EX 或 SE 阳性证据；当前事实层没有强标注 MA。\n【证据】dr_grade=1；MA present=template_only, source=grade_rule；HE/EX/SE present=false。\n【结论】该样本可按 Grade 1 的 MA-only 规则模板解释为轻度 DR，但不能表述为 RetSAM 检出 MA，也不能把 MA 当作 L3 可见病灶监督。",
        "{\"task\":\"L4_grade1_template\",\"dr_grade\":1,\"MA\":\"template_only\",\"ma_source\":\"grade_rule\",\"forbid\":\"RetSAM_detected_MA\"}",
    )
    cot_card(
        doc,
        "样例 4：L4 Grade 3 证据绑定分级",
        "aptos::all::0104b032c141；data/cropped/aptos/grade1_4/0104b032c141.png；grade=3，HE count=3，EX count=81，MA=unknown，SE=false。",
        "你是眼底分级助手。分级必须引用病灶证据；MA unknown 时不得编造 MA。",
        "<image> 请先核查可见病灶证据，再判断该 DR 分级是否有依据。",
        "【观察】可见少量暗红出血样病灶，同时可见大量亮黄色硬性渗出样病灶；未见可靠 SE 证据。\n【证据】HE present=true, count=3, area=157.0；EX present=true, count=81, area=30470.0；SE present=false；MA=unknown；dr_grade=3。\n【结论】监督分级为 DR Grade 3，主要解释证据是大量 EX 伴少量 HE。由于 MA unknown，不在解释中编造 MA。",
        "{\"task\":\"L4_evidence_bound_grading\",\"dr_grade\":3,\"evidence\":[\"EX\",\"HE\"],\"MA\":\"unknown\",\"source\":\"validated_clean\"}",
    )

    heading(doc, "六、训练流程设计", 1)
    para(
        doc,
        "训练不采用“每个小问题单独微调一次”的方式。每个小问题先生成独立 SFT 文件，便于抽检、采样和消融；真正训练时采用分阶段混合 continued SFT。"
        "每个阶段内部比例归一化为 100%，同一训练图像可以在不同阶段重复出现，但 train/val/test 必须先按图像级或 hash 级固定，防止泄露。"
    )
    para(
        doc,
        "固定事实层与数据隔离阶段。训练样本只从 validated_clean.jsonl 生成，先按 image_id、原始路径和图像 hash 做 group split，再生成 L2/L3/L4 问题。IDRiD test 保留为验证或测试，不进入训练；同一张图生成的多个问题必须留在同一个 split。",
        bold_prefix="Stage 0：",
    )
    para(
        doc,
        "视觉感知 SFT。从 Qwen3-VL-8B 基座开始，阶段内采用 L2 30% 与 L3 70%。此时不加入 DR grade，训练重点是让模型学会看视盘、黄斑、视杯、血管 QC，以及 MA/HE/EX/SE 的显式形态证据。",
        bold_prefix="Stage 1：",
    )
    para(
        doc,
        "证据绑定诊断 SFT。从 Stage 1 checkpoint 继续训练，阶段内采用 L2 20%、L3 50%、L4 30%。这一阶段开始加入 DR Grade 0-4 和 referable DR，但分级答案必须引用 L3 病灶证据；继续混入 L2/L3 是为了防止模型退化成只看标签分布的 grade classifier。",
        bold_prefix="Stage 2：",
    )
    para(
        doc,
        "平衡混合收敛阶段。使用更小学习率继续训练 0.5-1 个 epoch，阶段内采用 L2 20%、L3 55%、L4 25%。这一阶段主要稳定答案格式和证据链，同时避免高分级或少数类被过度上采样导致过诊断。",
        bold_prefix="Stage 3：",
    )
    para(
        doc,
        "可选的 DPO/MPO 修正阶段。只有当 Stage 2/3 后仍出现规则性幻觉时才使用，负样本来自真实错误点，例如把 template_only MA 写成 RetSAM 检出、把 cleaning_rule SE 写成阳性、vessel QC 失败时仍强判 A/V ratio 或 tortuosity。",
        bold_prefix="Stage 4：",
    )

    heading(doc, "七、评估与消融", 1)
    para(
        doc,
        "评估不只看最终 DR grade，而要同时报告 L2、L3、L4 分层结果和证据一致性。传统指标用于衡量关键医学任务，FunBench 风格分层评估用于定位模型短板，消融实验用于证明显式 CoT、强标注和清洗规则是否真正有效。"
    )
    table(
        doc,
        ["评估部分", "核心指标", "目的"],
        [
            ["L2 解剖", "laterality Accuracy、CDR MAE/Bucket Accuracy、Abstention Accuracy", "确认基础眼底结构感知和拒答能力"],
            ["L3 病灶", "MA/HE/EX/SE Macro-F1、Sensitivity、Specificity、AUPRC", "确认模型是否真的识别病灶，而非从分级猜测"],
            ["L4 分级", "Accuracy、Macro-F1、Quadratic Weighted Kappa、Referable DR Sens/Spec", "评估 DR 分级和转诊价值"],
            ["证据一致性", "Evidence Consistency Rate、Contradiction Rate、Rule Violation Rate", "检查解释是否引用正确病灶证据，是否编造 MA/SE/血管指标"],
            ["消融", "No Stage 1、No L2 replay、No L3 replay、No explicit lesion CoT、RetSAM raw vs clean、Grade-only L4", "验证分阶段训练、显式病灶描述和清洗规则的贡献"],
        ],
        widths=[3.3, 9.0, 8.5],
        font_size=8.4,
        long=True,
    )
    para(
        doc,
        "相似模型对比建议分三类进行：传统监督分类器（EfficientNet/ConvNeXt/ViT）只评 grade 和 referable DR；通用 MLLM（GPT-4o、Gemini、Claude、Qwen2.5-VL/Qwen3-VL 原始模型）使用同一 L2/L3/L4 prompt；本项目模型报告 Stage 1、Stage 2、Stage 3 以及可选 Stage 4 的逐阶段结果。",
    )

    heading(doc, "八、预期交付", 1)
    para(
        doc,
        "最终交付包括四部分：第一，统一可信事实层及统计文件；第二，L2/L3/L4 分任务 SFT 文件与阶段混合文件；第三，Qwen3-VL-8B LoRA/QLoRA 各阶段 checkpoint；第四，包含传统指标、FunBench 风格分层结果、消融结果和相似模型对比的评估报告。"
    )
    para(
        doc,
        "参考依据包括 FunBench: Benchmarking Fundus Reading Skills of MLLMs（MICCAI 2025）、LMOD（NAACL Findings 2025）、OphthaMMBench（2025）以及 2025 年关于 MLLM 诊断糖尿病视网膜病变的定量评估工作。",
        size=8.5,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
