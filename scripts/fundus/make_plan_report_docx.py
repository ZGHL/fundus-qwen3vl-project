from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "reports/qwen3vl_fundus_finetune_plan_report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        set_cell_shading(hdr[i], "2F5597")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def set_doc_style(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_doc():
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Qwen3-VL-8B 眼底图像微调方案汇报")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("基于 APTOS、DDR、IDRiD、FGADR 与 RetSAM 伪标签的分层 CoT/SFT 训练设计\n截至 2026-04-29").italic = True

    add_heading(doc, "一、项目目标", 1)
    p = doc.add_paragraph()
    p.add_run(
        "本方案目标是在 Qwen3-VL-8B 上构建眼底图像分层理解能力：先学习解剖结构和病灶证据，再学习基于证据的 DR 分级。"
        "核心思路不是直接训练模型输出等级，而是将任务拆成 L2 解剖感知、L3 病灶分析、L4 疾病诊断三个层级，降低标签捷径和病灶幻觉。"
    )

    add_heading(doc, "二、当前数据基础", 1)
    add_table(
        doc,
        ["数据集", "图像数", "RetSAM 有效数", "主要监督信息", "备注"],
        [
            ["APTOS", "3,662", "1,857", "DR grade + RetSAM HE/EX/SE", "RetSAM 仅覆盖 Grade 1-4"],
            ["DDR Grading", "5,006", "5,006", "DR grade + RetSAM HE/EX/SE", "Grade 5 已丢弃"],
            ["IDRiD train+test", "516", "413", "强 mask + grade + RetSAM 补充", "test 103 张不混入训练"],
            ["FGADR Seg-set", "1,842", "1,842", "强 mask + grade + RetSAM 补充", "强标注质量较高"],
            ["DDR lesion_seg", "757", "383", "强 lesion mask + RetSAM 补充", "valid/test 可用于 L3 泛化评估"],
        ],
        widths=[3.0, 2.0, 2.2, 5.0, 5.2],
    )
    add_bullets(
        doc,
        [
            "统一图像索引总数：11,783 张。",
            "RetSAM 有效 analysis-only 输出：9,501 张。",
            "强标注结构化样本：1,370 条，包括 IDRiD Stage1 Easy 与 FGADR lesion-only。",
            "清洗后可用样本：L2 9,501，L3 6,829，L4 9,493。",
        ],
    )

    add_heading(doc, "三、可信事实层与清洗规则", 1)
    p = doc.add_paragraph()
    p.add_run(
        "所有 SFT 样本只从 data/fundus_validated/validated_clean.jsonl 生成。"
        "强 mask 优先于 RetSAM，RetSAM 优先于 grade 规则模板；低置信、低 QC 或字段缺失时输出 unknown/拒答。"
    )
    add_table(
        doc,
        ["项目", "清洗/约束", "训练影响"],
        [
            ["MA", "RetSAM 不提供 MA；仅来自强 mask 或 Grade 1 规则模板", "禁止写成 RetSAM 检出 MA"],
            ["HE/EX", "RetSAM 需通过面积、数量和置信度过滤；强 mask 优先", "减少伪阳性病灶描述"],
            ["SE", "低置信或小面积 SE 降级为 false/unknown", "避免把噪声作为软性渗出监督"],
            ["eye_side/CDR", "OD 质量可靠时保留；坐标异常不自动否定 eye_side", "用于 L2 laterality 与 CDR"],
            ["血管指标", "vessel QC 失败时 A/V ratio、tortuosity 等置 unknown", "训练模型学会拒答"],
            ["分级一致性", "Grade 0/1 不允许编造 HE/EX/SE，高分级必须引用 L3 证据", "约束 L4 证据绑定诊断"],
        ],
        widths=[2.7, 7.6, 6.1],
    )

    add_heading(doc, "四、CoT 问题设计", 1)
    p = doc.add_paragraph()
    p.add_run(
        "参考 FunBench 的分层思想，每条训练样本只训练一个清晰能力点。"
        "system 放任务边界与拒答规则，user 只给观察任务不泄漏答案，assistant 给显式视觉描述、结构化证据、结论和最小 JSON。"
    )
    add_table(
        doc,
        ["层级", "子任务", "训练目标", "典型输出"],
        [
            ["L2 Anatomy", "左右眼、CDR、血管 QC", "识别视盘、黄斑、视杯和基础血管质量", "eye_side、cdr_bucket、unknown"],
            ["L3 Lesion", "MA/HE/EX/SE 显式识别、lesion-only、病灶负担", "学习颜色、形态、边界、数量、面积与病灶名映射", "lesions、count_bucket、area_bucket"],
            ["L4 Disease", "Grade 0/1 模板、Grade 2-4 证据绑定、冲突核查", "基于 L3 证据给出 DR 分级，避免只学 label prior", "dr_grade、evidence、needs_review"],
        ],
        widths=[3.0, 5.0, 6.0, 4.0],
    )

    add_heading(doc, "五、真实样本示例", 1)
    add_table(
        doc,
        ["样本 ID", "图像路径", "真实字段", "用于训练的重点"],
        [
            [
                "aptos::all::000c1434d8d7",
                "data/cropped/aptos/grade1_4/000c1434d8d7.png",
                "right eye；CDR=0.3495；HE=19；EX=26；SE 被清洗降级",
                "L2 laterality/CDR、L3 HE/EX、SE 拒答",
            ],
            [
                "idrid::train::IDRiD_001",
                "data/idrid/images/train/IDRiD_001.jpg",
                "MA=18；HE=15；EX area=12027；来源 strong_mask_stage1_easy",
                "强标注 L3 显式病灶 CoT",
            ],
            [
                "aptos::all::0024cdab0c1e",
                "data/cropped/aptos/grade1_4/0024cdab0c1e.png",
                "Grade 1；MA=template_only；HE/EX/SE=false",
                "L4 Grade 1 模板解释，禁止伪造 RetSAM MA",
            ],
            [
                "aptos::all::0104b032c141",
                "data/cropped/aptos/grade1_4/0104b032c141.png",
                "Grade 3；HE=3；EX=81；MA=unknown；SE=false",
                "L4 Grade 3 证据绑定分级",
            ],
        ],
        widths=[4.0, 5.5, 5.3, 4.0],
    )

    add_heading(doc, "六、训练流程", 1)
    p = doc.add_paragraph()
    p.add_run(
        "不为每个小问题单独微调模型，而是先生成分任务数据文件，再按阶段混合 continued SFT。"
        "每个阶段内部比例归一化为 100%；训练集样本可在不同阶段重复出现，但验证/测试必须按图像级隔离。"
    )
    add_table(
        doc,
        ["阶段", "数据比例", "训练目标", "说明"],
        [
            ["Stage 0", "-", "固定事实层与 split", "先按 image_id/路径/hash 分组，IDRiD test 不进训练"],
            ["Stage 0.5", "-", "生成子任务 SFT 文件", "L2/L3/L4 分文件，便于采样、抽检和消融"],
            ["Stage 1", "L2 30% + L3 70%", "视觉感知 SFT", "不加入 DR grade，先学解剖与病灶证据"],
            ["Stage 2", "L2 20% + L3 50% + L4 30%", "证据绑定诊断 SFT", "加入分级，但保留 L2/L3 replay 防遗忘"],
            ["Stage 3", "L2 20% + L3 55% + L4 25%", "小学习率混合收敛", "避免 grade 压过视觉证据"],
            ["Stage 4", "可选", "DPO/MPO 修正规则性错误", "重点处理 MA 伪造、SE 低置信、血管 QC 误判等问题"],
        ],
        widths=[2.4, 4.2, 4.2, 6.0],
    )
    add_bullets(
        doc,
        [
            "硬件：GB10 统一内存 128G，建议优先采用 Qwen3-VL-8B LoRA/QLoRA，而非一开始全量微调。",
            "Stage 1 重点建立看结构和病灶的能力；Stage 2 才加入 grade；Stage 3 用小学习率做最终混合收敛。",
            "NV 当前仅 27 条，更适合作为低权重“可见时识别、不可见不编造”监督，不适合大比例上采样。",
        ],
    )

    add_heading(doc, "七、防止数据泄露", 1)
    add_numbered(
        doc,
        [
            "先按图像级或 hash 级划分 train/val/test，再从每张图生成多个 L2/L3/L4 问题。",
            "同一张图生成的多个问题必须留在同一个 split，不能随机打散到训练和测试。",
            "IDRiD test 27 条强标注样本只用于验证或测试，不进入训练。",
            "清洗阈值和模板规则只根据训练集统计与人工抽检确定，不用测试集反复调参。",
            "跨数据集重复或近重复图按同一组处理，避免 APTOS/DDR/FGADR/IDRiD 间潜在重复泄露。",
        ],
    )

    add_heading(doc, "八、评估方案", 1)
    add_table(
        doc,
        ["评估维度", "关键指标", "主要目的"],
        [
            ["L2 解剖", "laterality Accuracy、CDR MAE/Bucket Accuracy、拒答准确率", "确认模型是否具备基础眼底结构感知"],
            ["L3 病灶", "MA/HE/EX/SE Macro-F1、Sensitivity、Specificity、AUPRC", "确认模型是否真的识别病灶，而非复述标签"],
            ["L4 分级", "Accuracy、Macro-F1、Quadratic Weighted Kappa、Referable DR Sens/Spec", "评价 DR 分级与转诊价值"],
            ["证据一致性", "Evidence Consistency Rate、Contradiction Rate、Rule Violation Rate", "检查分级解释是否引用正确病灶证据"],
            ["人工抽检", "Grade 0/1、MA unknown、SE cleaning_rule、vessel QC false", "发现自动指标难覆盖的幻觉和不当拒答"],
        ],
        widths=[3.3, 7.2, 6.2],
    )

    add_heading(doc, "九、消融与模型对比", 1)
    add_table(
        doc,
        ["实验", "对照设置", "观察重点"],
        [
            ["No Stage 1", "直接 L2/L3/L4 混合 SFT", "验证先学视觉证据是否必要"],
            ["No L2 replay", "Stage 2/3 不混入 L2", "检查 laterality/CDR 是否遗忘"],
            ["No L3 replay", "Stage 2/3 不混入 L3", "检查 L4 是否退化成 grade classifier"],
            ["No explicit lesion CoT", "L3 只保留 lesion-only 标签", "验证显式病灶描述是否提升 grounding"],
            ["RetSAM raw vs clean", "未清洗 RetSAM 对比清洗后 RetSAM", "验证清洗是否降低幻觉和冲突"],
            ["Grade-only L4", "答案只给 grade 不给证据", "验证证据绑定对一致性的贡献"],
        ],
        widths=[4.2, 6.2, 6.2],
    )
    add_table(
        doc,
        ["对比组", "模型/方法", "评价方式"],
        [
            ["传统监督基线", "EfficientNet / ConvNeXt / ViT DR classifier", "只评 grade、referable DR"],
            ["通用 MLLM", "GPT-4o、Gemini、Claude、Qwen2.5-VL/Qwen3-VL 原始模型", "同一 L2/L3/L4 prompt，评结构化答案与人工抽检"],
            ["医学/眼科 MLLM", "LLaVA-Med、Med-Flamingo、BiomedGPT 等可获得模型", "同一测试集、同一输出格式"],
            ["本项目模型", "Qwen3-VL-8B + LoRA/QLoRA", "报告 Stage 1/2/3/4 逐阶段结果"],
        ],
        widths=[3.2, 6.8, 6.6],
    )

    add_heading(doc, "十、预期交付", 1)
    add_bullets(
        doc,
        [
            "统一可信事实层：validated_clean.jsonl 及统计文件。",
            "分层 SFT 数据：L2/L3/L4 子任务文件与阶段混合文件。",
            "Qwen3-VL-8B LoRA/QLoRA checkpoint：Stage 1、Stage 2、Stage 3，可选 Stage 4。",
            "评估报告：传统指标、FunBench 风格分层结果、消融结果和相似模型对比。",
        ],
    )

    refs = doc.add_paragraph()
    refs.add_run("参考：").bold = True
    refs.add_run(
        "FunBench MICCAI 2025；LMOD NAACL Findings 2025；OphthaMMBench 2025；"
        "Can Multimodal Large Language Models Diagnose Diabetic Retinopathy from Fundus Photos? 2025。"
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
